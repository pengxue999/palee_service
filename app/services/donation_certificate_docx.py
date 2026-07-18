
from io import BytesIO

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.models.donation import Donation
from app.services.pdf.assets import IMAGE_DIR
from app.services.pdf.contexts.donation_certificate import (
    build_donation_certificate_context,
)

_LAO_FONT = "Noto Sans Lao"
_INK = RGBColor(0x11, 0x11, 0x11)


def _set_lao_font(run, *, size: float, bold: bool = False) -> None:
    run.font.name = _LAO_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = _INK
    # Ensure the complex-script (Lao) font is applied too, not just Latin.
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), _LAO_FONT)
    rfonts.set(qn("w:hAnsi"), _LAO_FONT)
    rfonts.set(qn("w:cs"), _LAO_FONT)
    # Word tracks bold separately for complex scripts.
    if bold:
        bcs = rpr.find(qn("w:bCs"))
        if bcs is None:
            bcs = rpr.makeelement(qn("w:bCs"), {})
            rpr.append(bcs)


def _add_line(
    document,
    text: str,
    *,
    size: float,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    space_after: float = 2,
    space_before: float = 0,
):
    paragraph = document.add_paragraph()
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(space_after)
    fmt.space_before = Pt(space_before)
    fmt.line_spacing = 1.3
    if text:
        run = paragraph.add_run(text)
        _set_lao_font(run, size=size, bold=bold)
    return paragraph


def _add_runs(paragraph, parts: list[tuple[str, bool]], *, size: float) -> None:
    for text, bold in parts:
        if not text:
            continue
        run = paragraph.add_run(text)
        _set_lao_font(run, size=size, bold=bold)


def build_donation_certificate_docx(donation: Donation) -> bytes:
    ctx = build_donation_certificate_context(donation)

    document = Document()

    # A4 landscape with matching margins (~10mm/14mm in the HTML).
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(2.6)
    section.right_margin = Cm(2.6)

    # --- Header (national heading, centered) ---
    _add_line(
        document,
        "ສາທາລະນະລັດ ປະຊາທິປະໄຕ ປະຊາຊົນລາວ",
        size=13,
        space_after=1,
    )
    _add_line(
        document,
        "ສັນຕິພາບ ເອກະລາດ ປະຊາທິປະໄຕ ເອກະພາບ ວັດທະນາຖາວອນ",
        size=11,
        space_after=1,
    )
    _add_line(document, "-------000-------", size=11, space_after=6)

    # --- Meta row: logo (left) + issue info (right) ---
    meta_table = document.add_table(rows=1, cols=2)
    meta_table.autofit = False
    logo_cell, info_cell = meta_table.rows[0].cells
    logo_cell.width = Cm(6)
    info_cell.width = Cm(17)

    logo_para = logo_cell.paragraphs[0]
    logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    logo_path = IMAGE_DIR / "logo.png"
    if logo_path.exists():
        logo_para.add_run().add_picture(str(logo_path), width=Cm(2.9))

    info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_runs(
        info_cell.paragraphs[0],
        [("ເລກທີ:............./ສບນກ", False)],
        size=11,
    )
    info_line2 = info_cell.add_paragraph()
    info_line2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _add_runs(
        info_line2,
        [(f"ນະຄອນຫຼວງວຽງຈັນ, ວັນທີ {ctx['issue_date']}", False)],
        size=11,
    )

    # --- Title + caption ---
    _add_line(
        document,
        str(ctx["certificate_title"]),
        size=18,
        bold=True,
        space_before=8,
        space_after=4,
    )
    _add_line(
        document,
        str(ctx["organization_caption"]),
        size=12,
        bold=True,
        space_after=10,
    )

    # --- Body content ---
    donor_section = str(ctx.get("donor_section") or "").strip()
    recipient = document.add_paragraph()
    recipient.alignment = WD_ALIGN_PARAGRAPH.LEFT
    recipient.paragraph_format.space_after = Pt(6)
    recipient.paragraph_format.line_spacing = 1.6
    recipient_parts: list[tuple[str, bool]] = [
        ("ມອບໃຫ້ແກ່: ", False),
        (str(ctx["donor_name"]), True),
    ]
    if donor_section:
        recipient_parts.append((f" ({donor_section})", False))
    _add_runs(recipient, recipient_parts, size=13)

    achievement = document.add_paragraph()
    achievement.alignment = WD_ALIGN_PARAGRAPH.LEFT
    achievement.paragraph_format.space_after = Pt(6)
    achievement.paragraph_format.line_spacing = 1.6
    _add_runs(
        achievement,
        [
            ("ທີ່ມີຜົນງານ: ປະກອບສ່ວນຊຸກຍູ້ສະໜັບສະໜູນ ", False),
            (str(ctx["donation_name"]), True),
            (f" {ctx['donation_amount_label']} ", False),
            (str(ctx["amount_text"]), True),
            (f" ໃຫ້{ctx['organization_name']}.", False),
        ],
        size=13,
    )

    body = document.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    body.paragraph_format.first_line_indent = Cm(1.8)
    body.paragraph_format.space_before = Pt(4)
    body.paragraph_format.line_spacing = 1.6
    _add_runs(
        body,
        [
            (
                f"ດັ່ງນັ້ນ, ຈຶ່ງໄດ້ມອບ{ctx['certificate_title']} "
                "ເພື່ອສະແດງຄວາມຮູ້ບຸນຄຸນ ແລະ "
                "ຈາລຶກເຖິງຄຸນງາມຄວາມດີຂອງທ່ານໄວ້ໃນປຶ້ມປະຫວັດສາດຂອງ "
                f"{ctx['organization_name']}.",
                False,
            )
        ],
        size=13,
    )

    # --- Signature area (right aligned) ---
    _add_line(
        document,
        "ຜູ້ອຳນວຍການ",
        size=13,
        bold=True,
        align=WD_ALIGN_PARAGRAPH.RIGHT,
        space_before=18,
        space_after=0,
    )
    # Blank space for the signature.
    _add_line(document, "", size=13, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=22)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
